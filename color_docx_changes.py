#!/usr/bin/env python3
"""
Color changed text in a revised .docx relative to an original .docx.

- Outputs a copy of the revised document with changed text colored.
- Does NOT generate Word Track Changes.
- Inserted paragraphs are colored entirely (no original counterpart).
- Replaced paragraphs: colors only the inserted/replaced spans at the word level.

Usage:
.venv313/Scripts/python.exe color_docx_changes.py --original "old.docx" --revised  "new.docx" --out "new_MARKED.docx" --rgb "CC0000"
  
"""

import argparse
import os
import re
import shutil
import tempfile
from copy import deepcopy
from difflib import SequenceMatcher
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor


def ensure_docx(path: str) -> tuple[str, bool]:
    """
    If *path* ends in .doc (old binary format), convert it to a temp .docx
    using Word via COM and return (tmp_docx_path, True).
    If it's already .docx (or .docx-compatible), return (path, False).
    The caller is responsible for deleting the temp file when created=True.
    """
    if path.lower().endswith(".docx"):
        return path, False

    # Try COM conversion (requires Microsoft Word installed)
    try:
        import pythoncom
        import win32com.client as win32
    except ImportError:
        raise ImportError(
            "The file is in legacy .doc format. Install pywin32 to auto-convert:\n"
            "  .venv313\\Scripts\\pip install pywin32\n"
            "Or manually save the file as .docx in Word first."
        )

    abs_path = os.path.abspath(path)
    tmp_dir = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + ".docx")

    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        # wdFormatXMLDocument = 12
        doc.SaveAs2(tmp_docx, FileFormat=12)
        doc.Close(False)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

    return tmp_docx, True


# ---------------------------
# Utilities
# ---------------------------

def normalize(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def is_empty_para(text: str) -> bool:
    return normalize(text) == ""

def tokenize_words(s: str) -> List[str]:
    """
    Tokenize into words + whitespace so we can re-assemble exactly.
    Example: "Hi there!" -> ["Hi", " ", "there", "!"]
    """
    if s is None:
        return []
    # split into word chars, whitespace, or single punctuation
    return re.findall(r"\w+|\s+|[^\w\s]", s, flags=re.UNICODE)

def hex_to_rgb(hex_str: str) -> RGBColor:
    hs = hex_str.strip().lstrip("#")
    if len(hs) != 6:
        raise ValueError("RGB hex must be 6 characters, e.g. CC0000")
    r = int(hs[0:2], 16)
    g = int(hs[2:4], 16)
    b = int(hs[4:6], 16)
    return RGBColor(r, g, b)

def copy_paragraph_style(dst_para, src_para):
    try:
        dst_para.style = src_para.style
    except Exception:
        pass
    try:
        dst_para.paragraph_format.alignment = src_para.paragraph_format.alignment
        dst_para.paragraph_format.left_indent = src_para.paragraph_format.left_indent
        dst_para.paragraph_format.right_indent = src_para.paragraph_format.right_indent
        dst_para.paragraph_format.first_line_indent = src_para.paragraph_format.first_line_indent
        dst_para.paragraph_format.space_before = src_para.paragraph_format.space_before
        dst_para.paragraph_format.space_after = src_para.paragraph_format.space_after
        dst_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing
    except Exception:
        pass

def extract_paragraph_texts(doc: Document) -> List[str]:
    return [p.text or "" for p in doc.paragraphs]

def build_compacted_index(texts: List[str]) -> Tuple[List[str], List[int]]:
    """
    Return (compacted_norm_texts, map_compacted_to_original_index),
    skipping empty paragraphs.
    """
    compacted = []
    mapping = []
    for i, t in enumerate(texts):
        if is_empty_para(t):
            continue
        compacted.append(normalize(t))
        mapping.append(i)
    return compacted, mapping


# ---------------------------
# Diff + Marking logic
# ---------------------------

def _set_run_color_inplace(r_elem, color: Optional[RGBColor]):
    """
    Set or remove the w:color element inside a <w:r>'s <w:rPr> in-place.
    color=None means remove any existing color override.
    """
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        if color is None:
            return
        rPr = OxmlElement('w:rPr')
        r_elem.insert(0, rPr)
    color_elem = rPr.find(qn('w:color'))
    if color is None:
        if color_elem is not None:
            rPr.remove(color_elem)
    else:
        if color_elem is None:
            color_elem = OxmlElement('w:color')
            rPr.append(color_elem)
        color_elem.set(qn('w:val'), str(color))


def _replace_run_text(r_elem, text: str):
    """Replace all <w:t> children in a <w:r> with a single new one."""
    for t in r_elem.findall(qn('w:t')):
        r_elem.remove(t)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)


def _split_run_inplace(
    r_elem,
    segments: List[tuple],  # (text: str, colorize: bool)
    change_color: RGBColor
):
    """
    Replace a single <w:r> element in-place with one <w:r> per segment,
    inserted at the same XML position (so surrounding non-run elements
    such as hyperlinks, bookmarks, footnote refs stay where they are).
    """
    parent = r_elem.getparent()
    for seg_text, colorize in segments:
        new_r = deepcopy(r_elem)
        _replace_run_text(new_r, seg_text)
        _set_run_color_inplace(new_r, change_color if colorize else None)
        r_elem.addprevious(new_r)
    parent.remove(r_elem)


def color_entire_paragraph(paragraph, color: RGBColor):
    """
    Color every run in the paragraph in-place, preserving all other formatting
    and the position of every non-run XML element (hyperlinks, bookmarks, etc.).
    """
    for run in paragraph.runs:
        _set_run_color_inplace(run._r, color)


def rewrite_paragraph_with_wordlevel_diff(
    paragraph,
    orig_text: str,
    rev_text: str,
    change_color: RGBColor
):
    """
    Color only the changed spans in a paragraph, modifying runs in-place.

    Key design: we never remove-then-append runs.  Instead we either
    modify a run's colour attribute directly (no split needed) or we
    deep-copy + split it at colour boundaries using addprevious(), then
    remove the original.  This guarantees that non-run XML siblings
    (hyperlinks, bookmarks, footnote/endnote refs, field codes …) are
    never displaced — solving the superscript-number reordering bug.
    """
    runs = list(paragraph.runs)
    runs_text = "".join(run.text or "" for run in runs)

    o = tokenize_words(orig_text)
    r = tokenize_words(runs_text)

    sm = SequenceMatcher(None, o, r, autojunk=False)
    ops = sm.get_opcodes()

    # Build character-level colour map over runs_text.
    token_starts: List[int] = []
    pos = 0
    for tok in r:
        token_starts.append(pos)
        pos += len(tok)
    token_starts.append(pos)  # sentinel

    colored_chars = bytearray(len(runs_text))  # 1 = should be coloured
    for tag, i1, i2, j1, j2 in ops:
        if tag != 'equal' and j2 > j1:
            cs = token_starts[j1]
            ce = token_starts[j2]
            for c in range(cs, min(ce, len(runs_text))):
                colored_chars[c] = 1

    # Process each run in-place.
    char_pos = 0
    for run in runs:
        text = run.text or ""
        n = len(text)
        if n == 0:
            continue

        # Compute colour segments for this run.
        segments: List[tuple] = []
        seg_start = 0
        while seg_start < n:
            abs_pos = char_pos + seg_start
            seg_col = colored_chars[abs_pos] if abs_pos < len(colored_chars) else 0
            seg_end = seg_start + 1
            while seg_end < n:
                abs_end = char_pos + seg_end
                c = colored_chars[abs_end] if abs_end < len(colored_chars) else 0
                if c != seg_col:
                    break
                seg_end += 1
            segments.append((text[seg_start:seg_end], bool(seg_col)))
            seg_start = seg_end

        if len(segments) == 1:
            # Entire run is one colour — modify in-place, no split.
            _, colorize = segments[0]
            if colorize:
                _set_run_color_inplace(run._r, change_color)
            # else: leave existing colour untouched
        else:
            # Run straddles a colour boundary — split in-place.
            _split_run_inplace(run._r, segments, change_color)

        char_pos += n

def mark_revised_document(
    original_path: str,
    revised_path: str,
    out_path: str,
    change_color_hex: str
):
    change_color = hex_to_rgb(change_color_hex)

    print(f"Loading documents...")
    orig_docx, orig_tmp = ensure_docx(original_path)
    rev_docx,  rev_tmp  = ensure_docx(revised_path)

    try:
        orig_doc = Document(orig_docx)
        rev_doc  = Document(rev_docx)
    finally:
        # clean up any temp dirs created during conversion
        if orig_tmp:
            shutil.rmtree(os.path.dirname(orig_docx), ignore_errors=True)
        if rev_tmp:
            shutil.rmtree(os.path.dirname(rev_docx), ignore_errors=True)

    orig_texts = extract_paragraph_texts(orig_doc)
    rev_texts = extract_paragraph_texts(rev_doc)

    print(f"  Original : {len(orig_texts)} paragraphs")
    print(f"  Revised  : {len(rev_texts)} paragraphs")

    # Build compacted sequences for paragraph-level alignment
    orig_seq, orig_map = build_compacted_index(orig_texts)
    rev_seq, rev_map = build_compacted_index(rev_texts)

    print("Diffing paragraphs...")
    sm = SequenceMatcher(None, orig_seq, rev_seq, autojunk=False)
    ops = sm.get_opcodes()

    n_inserted = n_replaced = n_skipped = 0

    # Walk paragraph-level ops and apply marking to the revised doc
    for tag, i1, i2, j1, j2 in ops:
        if tag == "equal":
            continue

        # indices in compacted space -> real paragraph indices
        orig_idxs = [orig_map[k] for k in range(i1, i2)]
        rev_idxs  = [rev_map[k]  for k in range(j1, j2)]

        if tag == "insert":
            # Paragraph(s) appear only in revised -> color whole paragraph(s)
            for ridx in rev_idxs:
                color_entire_paragraph(rev_doc.paragraphs[ridx], change_color)
                n_inserted += 1

        elif tag == "replace":
            # Try to pair up paragraphs one-to-one when counts match;
            # otherwise, color whole revised paragraphs in this block.
            if len(orig_idxs) == len(rev_idxs) and len(rev_idxs) > 0:
                for oidx, ridx in zip(orig_idxs, rev_idxs):
                    otext = orig_texts[oidx]
                    rtext = rev_texts[ridx]
                    # Skip formatting-only differences (text is identical)
                    if normalize(otext) == normalize(rtext):
                        n_skipped += 1
                        continue
                    rewrite_paragraph_with_wordlevel_diff(
                        rev_doc.paragraphs[ridx],
                        orig_text=otext,
                        rev_text=rtext,
                        change_color=change_color
                    )
                    n_replaced += 1
            else:
                for ridx in rev_idxs:
                    color_entire_paragraph(rev_doc.paragraphs[ridx], change_color)
                    n_replaced += 1

        elif tag == "delete":
            # Only in original, nothing to mark in revised
            pass

    print(f"  {n_inserted} inserted paragraph(s) colored entirely")
    print(f"  {n_replaced} changed paragraph(s) colored at word level")
    print(f"  {n_skipped} formatting-only paragraph(s) skipped")
    print(f"Saving → {out_path}")
    rev_doc.save(out_path)
    print("Done.")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--original", required=True, help="Path to original .docx")
    ap.add_argument("--revised", required=True, help="Path to revised .docx")
    ap.add_argument("--out", required=True, help="Output .docx path")
    ap.add_argument("--rgb", default="CC0000", help="Hex RGB for changed text (default CC0000)")
    args = ap.parse_args()

    mark_revised_document(args.original, args.revised, args.out, args.rgb)

if __name__ == "__main__":
    main()